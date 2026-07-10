"""Build a pandoc reference.docx from an md2docx theme.

Pandoc styles its .docx output from a *reference document*: it maps Markdown
constructs onto named styles ("Normal", "Heading 1", "Title", "Source Code",
"Block Text", ...). This script takes pandoc's default reference and restyles
those named styles to match an md2docx theme (colors, fonts, sizes, page
geometry), plus optionally bakes a footer / page numbers.

It reuses the typed Pydantic models from the Python implementation, so the same
theme YAML drives both the python-docx and pandoc outputs.

Usage:
    build_reference.py --base default.docx --config theme.yaml --output ref.docx
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Pt, RGBColor, Twips

from md2docx.config_models import Chrome, Theme, load_config_file
from md2docx.render.oxml import add_page_number_field, add_paragraph_borders, add_tab


def _style(doc, name):
    try:
        return doc.styles[name]
    except KeyError:
        return None


def _ensure_paragraph_style(doc, name):
    s = _style(doc, name)
    if s is None:
        s = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    return s


def _apply(style, *, name=None, size=None, color=None, bold=None, italic=None):
    if style is None:
        return
    f = style.font
    if name:
        f.name = name
    if size is not None:
        f.size = Pt(size)
    if color:
        f.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        f.bold = bold
    if italic is not None:
        f.italic = italic


def build(base: Path, theme: Theme, chrome: Chrome, output: Path) -> None:
    doc = Document(str(base))

    # Body / default text
    _apply(_style(doc, "Normal"), name=theme.font, size=theme.body.size, color=theme.colors.text)
    _apply(_style(doc, "Body Text"), name=theme.font, size=theme.body.size, color=theme.colors.text)
    _apply(_style(doc, "First Paragraph"), name=theme.font, size=theme.body.size, color=theme.colors.text)
    _apply(_style(doc, "Compact"), name=theme.font, size=theme.body.size, color=theme.colors.text)

    # Headings (h4-h6 reuse the h3 look)
    for level in (1, 2, 3):
        h = theme.headings.level(level)
        _apply(_style(doc, f"Heading {level}"), name=theme.font, size=h.size, color=h.color, bold=h.bold)
    h3 = theme.headings.level(3)
    for level in (4, 5, 6):
        _apply(_style(doc, f"Heading {level}"), name=theme.font, size=h3.size, color=h3.color, bold=h3.bold)

    # Title block (pandoc title/subtitle/author metadata)
    _apply(_style(doc, "Title"), name=theme.font, size=theme.cover.title_size, color=theme.colors.primary, bold=theme.cover.title_bold)
    _apply(_style(doc, "Subtitle"), name=theme.font, size=theme.cover.subtitle_size, color=theme.colors.muted)
    _apply(_style(doc, "Author"), name=theme.font, size=theme.cover.subtitle_size, color=theme.colors.muted)
    _apply(_style(doc, "Date"), name=theme.font, size=theme.cover.subtitle_size, color=theme.colors.muted)

    # Inline + block code
    _apply(_style(doc, "Verbatim Char"), name=theme.mono_font, size=theme.code.size, color=theme.colors.code_text)
    _apply(_ensure_paragraph_style(doc, "Source Code"), name=theme.mono_font, size=theme.code.size, color=theme.colors.code_text)

    # Blockquote + links
    _apply(_style(doc, "Block Text"), name=theme.font, size=theme.body.size, color=theme.colors.muted, italic=False)
    _apply(_style(doc, "Hyperlink"), color=theme.colors.secondary)

    # Page geometry
    section = doc.sections[0]
    section.page_width = Twips(theme.page.width)
    section.page_height = Twips(theme.page.height)
    section.top_margin = Twips(theme.page.margin.top)
    section.bottom_margin = Twips(theme.page.margin.bottom)
    section.left_margin = Twips(theme.page.margin.left)
    section.right_margin = Twips(theme.page.margin.right)

    # Footer (baked into the reference; pandoc keeps the reference's footers)
    if chrome.footer_text or chrome.page_numbers:
        section.footer.is_linked_to_previous = False
        para = section.footer.paragraphs[0]
        para.paragraph_format.tab_stops.add_tab_stop(Twips(theme.page.content_width), WD_TAB_ALIGNMENT.RIGHT)
        add_paragraph_borders(para, edges={"top": (4, theme.colors.divider, 6)})
        if chrome.footer_text:
            run = para.add_run(chrome.footer_text)
            run.font.name = theme.font
            run.font.size = Pt(theme.footer.size)
            run.font.color.rgb = RGBColor.from_string(theme.colors.muted)
        if chrome.page_numbers:
            add_tab(para, theme.font)
            label = para.add_run("Page ")
            label.font.name = theme.font
            label.font.size = Pt(theme.footer.size)
            label.font.color.rgb = RGBColor.from_string(theme.colors.muted)
            add_page_number_field(
                para, color=theme.colors.muted, font=theme.font, size_half_points=int(theme.footer.size * 2)
            )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


def main() -> None:
    ap = argparse.ArgumentParser(description="Build a pandoc reference.docx from an md2docx theme.")
    ap.add_argument("--base", required=True, type=Path, help="pandoc default reference.docx")
    ap.add_argument("--config", required=True, type=Path, help="theme/chrome YAML or JSON")
    ap.add_argument("--output", required=True, type=Path, help="reference.docx to write")
    args = ap.parse_args()

    data = load_config_file(args.config)
    theme = Theme.model_validate(data.get("theme", {}))
    chrome = Chrome.model_validate(data.get("chrome", {}))
    build(args.base, theme, chrome, args.output)
    print(f"wrote {args.output}")


if __name__ == "__main__":
    main()
