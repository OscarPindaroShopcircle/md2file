"""Walk a markdown-it ``inline`` token's children into styled python-docx runs.

With ``html=False`` markdown-it emits raw HTML as literal text, so the
strip-and-warn policy (drop tags, keep inner text, ``<br>`` -> line break, warn)
lives in :func:`_emit_text`.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx.image.image import Image as DocxImage
from docx.shared import Emu, Inches, Pt, RGBColor

from .context import RenderContext, RunStyle, base_style
from .oxml import add_hyperlink, run_in, set_char_spacing, set_run_shading

_TAG_RE = re.compile(r"<\/?\s*([a-zA-Z][a-zA-Z0-9]*)(?:\s[^>]*?)?\/?\s*>")


def render_inline(paragraph, token, ctx: RenderContext, base: RunStyle | None = None) -> None:
    """Append runs for ``token`` (an inline token) to ``paragraph``."""
    base = base or base_style(ctx.theme)
    theme = ctx.theme
    children = token.children or []

    state = {"bold": 1 if base.bold else 0, "italic": 1 if base.italic else 0, "link": None}

    def make_run(text: str = "", *, mono: bool = False):
        link = state["link"]
        run = run_in(link, paragraph) if link is not None else paragraph.add_run()
        run.text = text
        run.font.name = theme.mono_font if mono else base.font
        run.font.size = Pt(base.size)
        if mono:
            color = theme.colors.secondary if link is not None else theme.colors.code_text
        else:
            color = theme.colors.secondary if link is not None else base.color
        run.font.color.rgb = RGBColor.from_string(color)
        run.bold = True if state["bold"] > 0 else None
        run.italic = True if state["italic"] > 0 else None
        if link is not None:
            run.underline = True
        if mono:
            set_run_shading(run, theme.colors.code_bg)
        return run

    def line_break():
        run = run_in(state["link"], paragraph) if state["link"] is not None else paragraph.add_run()
        run.add_break()

    for child in children:
        t = child.type
        if t == "text":
            _emit_text(child.content, make_run, line_break, ctx)
        elif t == "strong_open":
            state["bold"] += 1
        elif t == "strong_close":
            state["bold"] -= 1
        elif t == "em_open":
            state["italic"] += 1
        elif t == "em_close":
            state["italic"] -= 1
        elif t in ("s_open", "s_close"):
            continue
        elif t == "code_inline":
            make_run(child.content, mono=True)
        elif t == "softbreak":
            make_run(" ")
        elif t == "hardbreak":
            line_break()
        elif t == "link_open":
            state["link"] = add_hyperlink(paragraph, child.attrGet("href") or "")
        elif t == "link_close":
            state["link"] = None
        elif t == "image":
            _add_image(paragraph, child, ctx, base)
        elif t == "html_inline":
            _emit_text(child.content, make_run, line_break, ctx)
        elif child.content:
            make_run(child.content)


def _emit_text(text: str, make_run, line_break, ctx: RenderContext) -> None:
    """Emit a text string, stripping literal HTML tags (keep inner text,
    ``<br>`` -> break, warn on other tags)."""
    if "<" not in text:
        if text:
            make_run(text)
        return
    last = 0
    for m in _TAG_RE.finditer(text):
        if m.start() > last:
            make_run(text[last : m.start()])
        tag = m.group(1).lower()
        if tag == "br":
            line_break()
        else:
            ctx.warn(tag)
        last = m.end()
    if last < len(text):
        make_run(text[last:])


def _add_image(paragraph, token, ctx: RenderContext, base: RunStyle) -> None:
    src = token.attrGet("src") or ""
    alt = token.content or src
    path = Path(src) if Path(src).is_absolute() else (ctx.doc_dir / src)
    try:
        max_emu = Inches(ctx.theme.image.max_width_px / 96)
        with path.open("rb") as fh:
            info = DocxImage.from_blob(fh.read())
        dpi = info.horz_dpi or 96
        natural = Inches(info.px_width / dpi)
        width = Emu(min(int(natural), int(max_emu)))
        run = paragraph.add_run()
        run.add_picture(str(path), width=width)
    except Exception:
        ctx.warn(f"image-not-found:{src}")
        run = paragraph.add_run(f"[image: {alt}]")
        run.italic = True
        run.font.name = base.font
        run.font.size = Pt(base.size)
        run.font.color.rgb = RGBColor.from_string(ctx.theme.colors.muted)
