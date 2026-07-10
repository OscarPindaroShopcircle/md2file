"""Low-level python-docx / OOXML helpers.

python-docx has no public API for several things we need (paragraph & cell
shading, arbitrary borders, hyperlinks, list numbering, page-number fields), so
these build the raw ``w:...`` elements. Everything OOXML-poking lives here.
"""

from __future__ import annotations

from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.text.run import Run


def _edge(tag: str, size: int, color: str, space: int = 0) -> OxmlElement:
    el = OxmlElement(f"w:{tag}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(size))
    el.set(qn("w:space"), str(space))
    el.set(qn("w:color"), color)
    return el


def add_paragraph_borders(paragraph, *, edges: dict[str, tuple[int, str, int]]) -> None:
    """Add borders to a paragraph. ``edges`` maps 'top'/'bottom'/'left'/'right'
    to (size, color, space)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for name, (size, color, space) in edges.items():
        pBdr.append(_edge(name, size, color, space))
    pPr.append(pBdr)


def set_paragraph_shading(paragraph, fill: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def set_run_shading(run: Run, fill: str) -> None:
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    rPr.append(shd)


def set_char_spacing(run: Run, twentieths: int) -> None:
    """Letter spacing, in twentieths of a point."""
    rPr = run._r.get_or_add_rPr()
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:val"), str(twentieths))
    rPr.append(spc)


def set_table_borders(table, size: int, color: str) -> None:
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borders.append(_edge(name, size, color))
    tblPr.append(borders)


def set_cell_shading(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_hyperlink(paragraph, url: str):
    """Append an empty ``w:hyperlink`` bound to ``url`` and return the element.
    Add runs into it with :func:`run_in`."""
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    paragraph._p.append(hyperlink)
    return hyperlink


def run_in(parent, paragraph) -> Run:
    """Create a ``w:r`` inside an arbitrary parent element (e.g. a hyperlink) and
    wrap it as a python-docx Run for normal styling."""
    r = OxmlElement("w:r")
    parent.append(r)
    return Run(r, paragraph)


def apply_numbering(paragraph, num_id: int, ilvl: int) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    il = OxmlElement("w:ilvl")
    il.set(qn("w:val"), str(ilvl))
    ni = OxmlElement("w:numId")
    ni.set(qn("w:val"), str(num_id))
    numPr.append(il)
    numPr.append(ni)
    pPr.append(numPr)


def add_page_number_field(paragraph, *, color: str, font: str, size_half_points: int) -> None:
    """Append a PAGE field (current page number) to ``paragraph``."""
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rPr.append(rfonts)
    col = OxmlElement("w:color")
    col.set(qn("w:val"), color)
    rPr.append(col)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size_half_points))
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


def add_tab(paragraph, run_font: str | None = None) -> Run:
    run = paragraph.add_run()
    run._r.append(OxmlElement("w:tab"))
    if run_font:
        run.font.name = run_font
    return run


# Re-exported for the numbering module.
__all__ = [
    "OxmlElement",
    "parse_xml",
    "nsdecls",
    "qn",
    "add_paragraph_borders",
    "set_paragraph_shading",
    "set_run_shading",
    "set_char_spacing",
    "set_table_borders",
    "set_cell_shading",
    "add_hyperlink",
    "run_in",
    "apply_numbering",
    "add_page_number_field",
    "add_tab",
]
