"""Block-level builders: create configured python-docx elements from tokens,
parameterized entirely by the theme."""

from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Pt, RGBColor, Twips

from .context import RenderContext, RunStyle, base_style
from .inline import render_inline
from .oxml import (
    add_page_number_field,
    add_paragraph_borders,
    add_tab,
    set_cell_shading,
    set_char_spacing,
    set_paragraph_shading,
    set_table_borders,
)

if TYPE_CHECKING:
    from ..config_models import Chrome, Theme

_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


def configure_document(doc, theme: "Theme") -> None:
    normal = doc.styles["Normal"]
    normal.font.name = theme.font
    normal.font.size = Pt(theme.body.size)
    normal.font.color.rgb = RGBColor.from_string(theme.colors.text)

    section = doc.sections[0]
    section.page_width = Twips(theme.page.width)
    section.page_height = Twips(theme.page.height)
    section.top_margin = Twips(theme.page.margin.top)
    section.bottom_margin = Twips(theme.page.margin.bottom)
    section.left_margin = Twips(theme.page.margin.left)
    section.right_margin = Twips(theme.page.margin.right)


def heading(doc, level: int, inline_token, ctx: RenderContext):
    theme = ctx.theme
    h = theme.headings.level(level)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(h.before)
    p.paragraph_format.space_after = Pt(h.after)
    p.paragraph_format.line_spacing = theme.body.line
    p.paragraph_format.keep_with_next = True
    render_inline(p, inline_token, ctx, RunStyle(color=h.color, font=theme.font, size=h.size, bold=h.bold))
    if h.rule:
        add_paragraph_borders(p, edges={"bottom": (h.rule_size, theme.colors.divider, 4)})
    return p


def body(doc, inline_token, ctx: RenderContext, *, quote: bool = False, center: bool = False):
    theme = ctx.theme
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(theme.body.after)
    p.paragraph_format.line_spacing = theme.body.line
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base = base_style(theme)
    if quote:
        base = RunStyle(color=theme.colors.muted, font=theme.font, size=theme.body.size)
        p.paragraph_format.left_indent = Twips(360)
        add_paragraph_borders(p, edges={"left": (18, theme.colors.secondary, 12)})
    render_inline(p, inline_token, ctx, base)
    return p


def list_item(doc, inline_token, ctx: RenderContext, num_id: int, ilvl: int):
    theme = ctx.theme
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(theme.body.list_after)
    p.paragraph_format.line_spacing = theme.body.line
    render_inline(p, inline_token, ctx, base_style(theme))
    from .oxml import apply_numbering

    apply_numbering(p, num_id, ilvl)
    return p


def divider(doc, ctx: RenderContext):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    add_paragraph_borders(p, edges={"bottom": (6, ctx.theme.colors.divider, 1)})
    return p


def code_block(doc, code: str, ctx: RenderContext):
    theme = ctx.theme
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.2
    lines = code.rstrip("\n").split("\n")
    for i, ln in enumerate(lines):
        run = p.add_run()
        if i > 0:
            run.add_break()
        run.add_text(ln if ln else " ")
        run.font.name = theme.mono_font
        run.font.size = Pt(theme.code.size)
        run.font.color.rgb = RGBColor.from_string(theme.colors.code_text)
    set_paragraph_shading(p, theme.colors.code_bg)
    add_paragraph_borders(
        p,
        edges={e: (2, theme.colors.divider, 6) for e in ("top", "bottom", "left", "right")},
    )
    return p


def table(doc, header_cells: list, rows: list[list], aligns: list[str], ctx: RenderContext):
    """header_cells / rows hold inline tokens (or None). aligns is per column."""
    theme = ctx.theme
    ncol = len(header_cells) or (len(rows[0]) if rows else 1)
    tbl = doc.add_table(rows=0, cols=ncol)
    tbl.autofit = False
    col_w = Twips(theme.page.content_width // ncol)

    def fill(cell, inline_token, align, header):
        cell.width = col_w
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = theme.body.line
        para.alignment = _ALIGN.get(align, WD_ALIGN_PARAGRAPH.LEFT)
        if header:
            set_cell_shading(cell, theme.table.header_fill)
            base = RunStyle(color=theme.table.header_color, font=theme.font, size=theme.body.size, bold=True)
        else:
            base = base_style(theme)
        if inline_token is not None:
            render_inline(para, inline_token, ctx, base)

    if header_cells:
        cells = tbl.add_row().cells
        for i, tok in enumerate(header_cells):
            fill(cells[i], tok, aligns[i] if i < len(aligns) else "left", True)
    for row in rows:
        cells = tbl.add_row().cells
        for i, tok in enumerate(row):
            fill(cells[i], tok, aligns[i] if i < len(aligns) else "left", False)

    set_table_borders(tbl, theme.table.border_size, theme.table.border_color)
    return tbl


# --- chrome --------------------------------------------------------------


def cover(doc, chrome: "Chrome", theme: "Theme", ctx: RenderContext) -> None:
    logo_path = Path(chrome.logo) if chrome.logo else None
    if logo_path and logo_path.exists():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(18)
        try:
            from docx.shared import Inches

            p.add_run().add_picture(str(logo_path), width=Inches(chrome.logo_width / 96))
        except Exception:
            ctx.warn(f"logo-not-usable:{chrome.logo}")

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(theme.cover.top_space)

    if chrome.eyebrow:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(chrome.eyebrow.upper())
        run.bold = True
        run.font.name = theme.font
        run.font.size = Pt(theme.cover.eyebrow_size)
        run.font.color.rgb = RGBColor.from_string(theme.colors.secondary)
        set_char_spacing(run, 40)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(chrome.title or "")
    run.bold = theme.cover.title_bold
    run.font.name = theme.font
    run.font.size = Pt(theme.cover.title_size)
    run.font.color.rgb = RGBColor.from_string(theme.colors.primary)

    for sub in chrome.subtitles:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(sub)
        run.font.name = theme.font
        run.font.size = Pt(theme.cover.subtitle_size)
        run.font.color.rgb = RGBColor.from_string(theme.colors.muted)

    divider(doc, ctx)
    doc.add_page_break()


def footer(doc, theme: "Theme", footer_text: str | None, page_numbers: bool) -> None:
    section = doc.sections[0]
    section.footer.is_linked_to_previous = False
    para = section.footer.paragraphs[0]
    para.paragraph_format.tab_stops.add_tab_stop(Twips(theme.page.content_width), WD_TAB_ALIGNMENT.RIGHT)
    add_paragraph_borders(para, edges={"top": (4, theme.colors.divider, 6)})

    if footer_text:
        run = para.add_run(footer_text)
        run.font.name = theme.font
        run.font.size = Pt(theme.footer.size)
        run.font.color.rgb = RGBColor.from_string(theme.colors.muted)
    if page_numbers:
        add_tab(para, theme.font)
        label = para.add_run("Page ")
        label.font.name = theme.font
        label.font.size = Pt(theme.footer.size)
        label.font.color.rgb = RGBColor.from_string(theme.colors.muted)
        add_page_number_field(
            para, color=theme.colors.muted, font=theme.font, size_half_points=int(theme.footer.size * 2)
        )
