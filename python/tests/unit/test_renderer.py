"""Unit tests for the block renderer's state handling (lists, tables, etc.)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from md2docx.render.context import RenderContext
from md2docx.render.numbering import NumberingManager
from md2docx.render.renderer import render


def _render(md, theme, text):
    doc = Document()
    ctx = RenderContext(theme=theme, doc_dir=Path("."))
    ctx.numbering = NumberingManager(doc, theme)
    render(doc, md.parse(text), ctx)
    return doc, ctx


def _numpr(paragraph):
    numpr = paragraph._p.find(qn("w:pPr") + "/" + qn("w:numPr"))
    if numpr is None:
        return None
    ilvl = numpr.find(qn("w:ilvl")).get(qn("w:val"))
    num_id = numpr.find(qn("w:numId")).get(qn("w:val"))
    return int(ilvl), int(num_id)


def test_nested_bullets_track_level(md, theme):
    doc, _ = _render(md, theme, "- a\n- b\n  - nested\n- c")
    listed = [p for p in doc.paragraphs if _numpr(p) is not None]
    levels = [_numpr(p)[0] for p in listed]
    assert 0 in levels and 1 in levels  # top-level and nested


def test_ordered_lists_get_distinct_num_ids(md, theme):
    doc, _ = _render(md, theme, "1. a\n2. b\n\ntext\n\n1. x\n2. y")
    num_ids = {_numpr(p)[1] for p in doc.paragraphs if _numpr(p) is not None}
    # two separate ordered lists -> two numbering instances (so each restarts)
    assert len(num_ids) == 2


def test_table_structure_and_header_bold(md, theme):
    doc, _ = _render(md, theme, "| A | B |\n|:--|--:|\n| 1 | 2 |")
    assert len(doc.tables) == 1
    tbl = doc.tables[0]
    assert tbl.rows[0].cells[0].text == "A"
    assert tbl.rows[1].cells[1].text == "2"
    header_run = tbl.rows[0].cells[0].paragraphs[0].runs[0]
    assert header_run.bold is True


def test_code_blockquote_divider_present(md, theme):
    doc, _ = _render(md, theme, "> quote\n\n```\ncode\n```\n\n---\n")
    body = "\n".join(p.text for p in doc.paragraphs)
    assert "quote" in body and "code" in body
    # divider is an empty paragraph carrying a bottom border
    assert any(p._p.find(qn("w:pPr") + "/" + qn("w:pBdr")) is not None for p in doc.paragraphs)
